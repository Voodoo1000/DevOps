from rest_framework.viewsets import GenericViewSet
from rest_framework import mixins, serializers
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.authentication import BasicAuthentication
from app.middlewares import CsrfExemptSessionAuthentication
from studentDormitory.models import Student, Room, DutySchedule, Staff, RepairRequests
from studentDormitory.serializers import StudentSerializer, RoomSerializer, DutyScheduleSerializer, StaffSerializer, RepairRequestsSerializer
from django.db.models import Avg, Count, Max, Min
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import  User
from django.http import FileResponse
from openpyxl import Workbook
from docx import Document 
import io
from django.utils.decorators import method_decorator

class StudentViewset(
	mixins.CreateModelMixin,
	mixins.UpdateModelMixin,
	mixins.RetrieveModelMixin,
	mixins.DestroyModelMixin,
	mixins.ListModelMixin,
	GenericViewSet):
	queryset = Student.objects.all()
	serializer_class = StudentSerializer
	authentication_classes =  (CsrfExemptSessionAuthentication, BasicAuthentication)
	def get_queryset(self):
		qs = super().get_queryset()
		
		if self.request.user.is_superuser:
				return qs
		
		return qs.filter(user=self.request.user)
	
	class StatsSerializer(serializers.Serializer):
		count = serializers.IntegerField()
		avg = serializers.FloatField()
		max = serializers.IntegerField()
		min = serializers.IntegerField()

	@action(detail=False, methods=["GET"], url_path="stats")
	def get_stats(self, *args, **kwargs):
		stats = Student.objects.aggregate(
			count = Count("*"),
			avg = Avg("id"),
			max = Max("id"),
			min = Min("id"),
		)

		serializer = self.StatsSerializer(instance = stats)

		return Response(serializer.data)

	@action(detail=False, methods=["GET"], url_path="export-excel")
	def export_to_excel(self, request, *args, **kwargs):
		workbook = Workbook()
		sheet = workbook.active
		sheet.title = "Students"
		sheet.append(["ID", "ФИО", "Группа", "Номер комнаты"])
		for student in Student.objects.select_related('room').all():
			sheet.append([student.id, student.name, student.group, student.room.number])

		file_stream = io.BytesIO()
		workbook.save(file_stream)
		file_stream.seek(0)

		response = FileResponse(file_stream, as_attachment=True, filename="students.xlsx")
		return response

	@action(detail=False, methods=["GET"], url_path="export-word")
	def export_to_word(self, request, *args, **kwargs):
		document = Document()
		document.add_heading("Students List", level=1)

		# Данные
		for student in Student.objects.select_related('room').all():
				document.add_paragraph(f"ID: {student.id}, ФИО: {student.name}, Группа: {student.group}, Номер комнаты: {student.room.number}")

		# Создаем файл в памяти
		file_stream = io.BytesIO()
		document.save(file_stream)
		file_stream.seek(0)

		response = FileResponse(file_stream, as_attachment=True, filename="students.docx")
		return response


class RoomViewset(
	mixins.CreateModelMixin,
	mixins.UpdateModelMixin,
	mixins.RetrieveModelMixin,
	mixins.DestroyModelMixin,
	mixins.ListModelMixin,
	GenericViewSet):
	queryset = Room.objects.all()
	serializer_class = RoomSerializer
	authentication_classes =  (CsrfExemptSessionAuthentication, BasicAuthentication)
	def get_queryset(self):
		qs = super().get_queryset()
		
		if self.request.user.is_superuser:
				return qs
		
		return qs.filter(user=self.request.user)
	
	class StatsSerializer(serializers.Serializer):
		count = serializers.IntegerField()
		avg = serializers.FloatField()
		max = serializers.IntegerField()
		min = serializers.IntegerField()

	@action(detail=False, methods=["GET"], url_path="stats")
	def get_stats(self, *args, **kwargs):
		stats = Room.objects.aggregate(
			count = Count("*"),
			avg = Avg("id"),
			max = Max("id"),
			min = Min("id"),
		)

		serializer = self.StatsSerializer(instance = stats)

		return Response(serializer.data)

class DutyScheduleViewset(
	mixins.CreateModelMixin,
	mixins.UpdateModelMixin,
	mixins.RetrieveModelMixin,
	mixins.DestroyModelMixin,
	mixins.ListModelMixin,
	GenericViewSet):
	queryset = DutySchedule.objects.all()
	serializer_class = DutyScheduleSerializer
	authentication_classes =  (CsrfExemptSessionAuthentication, BasicAuthentication)
	def get_queryset(self):
		qs = super().get_queryset()
		
		if self.request.user.is_superuser:
				return qs
		
		return qs.filter(user=self.request.user)
	
	class StatsSerializer(serializers.Serializer):
		count = serializers.IntegerField()
		avg = serializers.FloatField()
		max = serializers.IntegerField()
		min = serializers.IntegerField()

	@action(detail=False, methods=["GET"], url_path="stats")
	def get_stats(self, *args, **kwargs):
		stats = DutySchedule.objects.aggregate(
			count = Count("*"),
			avg = Avg("id"),
			max = Max("id"),
			min = Min("id"),
		)

		serializer = self.StatsSerializer(instance = stats)

		return Response(serializer.data)

class StaffViewset(
	mixins.CreateModelMixin,
	mixins.UpdateModelMixin,
	mixins.RetrieveModelMixin,
	mixins.DestroyModelMixin,
	mixins.ListModelMixin,
	GenericViewSet):
	queryset = Staff.objects.all()
	serializer_class = StaffSerializer
	authentication_classes =  (CsrfExemptSessionAuthentication, BasicAuthentication)
	def get_queryset(self):
		qs = super().get_queryset()
		
		if self.request.user.is_superuser:
				return qs
		
		return qs.filter(user=self.request.user)
	
	class StatsSerializer(serializers.Serializer):
		count = serializers.IntegerField()
		avg = serializers.FloatField()
		max = serializers.IntegerField()
		min = serializers.IntegerField()

	@action(detail=False, methods=["GET"], url_path="stats")
	def get_stats(self, *args, **kwargs):
		stats = Staff.objects.aggregate(
			count = Count("*"),
			avg = Avg("id"),
			max = Max("id"),
			min = Min("id"),
		)

		serializer = self.StatsSerializer(instance = stats)

		return Response(serializer.data)

class RepairRequestsViewset(
	mixins.CreateModelMixin,
	mixins.UpdateModelMixin,
	mixins.RetrieveModelMixin,
	mixins.DestroyModelMixin,
	mixins.ListModelMixin,
	GenericViewSet):
	queryset = RepairRequests.objects.all()
	serializer_class = RepairRequestsSerializer
	authentication_classes =  (CsrfExemptSessionAuthentication, BasicAuthentication)
	def get_queryset(self):
		qs = super().get_queryset()
		
		if self.request.user.is_superuser:
				return qs
		
		return qs.filter(user=self.request.user)
	
	class StatsSerializer(serializers.Serializer):
		count = serializers.IntegerField()
		avg = serializers.FloatField()
		max = serializers.IntegerField()
		min = serializers.IntegerField()

	@action(detail=False, methods=["GET"], url_path="stats")
	def get_stats(self, *args, **kwargs):
		stats = RepairRequests.objects.aggregate(
			count = Count("*"),
			avg = Avg("id"),
			max = Max("id"),
			min = Min("id"),
		)

		serializer = self.StatsSerializer(instance = stats)

		return Response(serializer.data)
	
	
class UserViewset(GenericViewSet):
	@action(url_path="info", methods=["GET"], detail=False)
	def get_info(self, request,  *args, **kwargs):
		data={
			"is_authenticated": request.user.is_authenticated
		}
		if request.user.is_authenticated:
			data.update({
				"username": request.user.username,
				"user_id": request.user.id,
				"is_superuser": request.user.is_superuser
			})

		return Response(data)

	@action(url_path="list", methods=["GET"], detail=False)
	def list_users(self, request, *args, **kwargs):
		if not request.user.is_superuser:
			return Response({"error": "Forbidden"}, status=403)
        
		users = User.objects.values("id", "username")
		return Response(users)
	
	@method_decorator(csrf_exempt)
	@action(url_path="login", methods=["POST"], detail=False)
	def login(self, request, *args, **kwargs):
		username = request.data.get("user")
		password = request.data.get("password")

		user = authenticate(request, username=username, password=password)
		if user is not None:
				login(request, user)
				return Response({"success": True})
		else:
				return Response({"error": "Invalid credentials"}, status=400)

	@action(url_path="logout", methods=["POST"], detail=False)
	def logout(self, request, *args, **kwargs):
			logout(request)
			return Response({"success": True})